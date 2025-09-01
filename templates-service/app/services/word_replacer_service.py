from docx import Document
from io import BytesIO
import base64
from typing import Union, Dict, Any, List, Tuple
from exceptions.tributarios_exception import TributarioException

class WordReplacerService:

    def _build_mapping(self, metadata: Dict[str, Any]) -> Dict[str, str]:
        """
        {"demandado_nombre": "Ramses"} -> {"{{DEMANDADO_NOMBRE}}": "Ramses"}
        """
        mapping = {}
        for k, v in (metadata or {}).items():
            if v is None:
                continue
            mapping[f"{{{{{str(k).upper()}}}}}"] = str(v)
        return mapping

    # --- Reemplazo por run, incluyendo placeholders cortados en varios runs ---
    def _replace_placeholders_in_runs(self, runs, mapping: Dict[str, str]):
        """
        Reemplaza placeholders del tipo {{KEY}} tanto si están en un único run
        como si aparecen repartidos en múltiples runs (p.ej. '{{' | 'KEY' | '}}').
        Conserva el formato del resto del párrafo. El texto reemplazado quedará
        con el formato del primer run que contenía '{{'.
        """
        i = 0
        while i < len(runs):
            text_i = runs[i].text or ""

            # 1) Caso rápido: placeholder completo dentro del mismo run
            changed = False
            for ph, val in mapping.items():
                if ph in text_i:
                    text_i = text_i.replace(ph, val)
                    changed = True
            if changed:
                runs[i].text = text_i
                i += 1
                continue

            # 2) Caso partido: buscamos '{{' y acumulamos hasta '}}' cruzando runs
            pos = text_i.find('{{')
            if pos == -1:
                i += 1
                continue

            # Podría haber varios '{{' en el mismo run; procesamos uno por iteración
            collected = text_i[pos:]  # desde '{{' de este run
            consumed: List[Tuple[int, int, int]] = [(i, pos, len(text_i))]  # (idx_run, from,to)
            j = i + 1
            while j < len(runs) and '}}' not in collected:
                t = runs[j].text or ""
                collected += t
                consumed.append((j, 0, len(t)))
                j += 1

            if '}}' not in collected:
                # No se cerró; no tocamos nada y avanzamos
                i += 1
                continue

            # Extraemos primer token {{...}}
            open_idx = collected.find('{{')
            close_idx = collected.find('}}', open_idx + 2)
            token = collected[open_idx:close_idx + 2]  # incluye '}}'

            # Buscamos en mapping (tolerancia por mayúsc/minúsc internas)
            replacement = mapping.get(token)
            if replacement is None:
                middle = token[2:-2].upper()
                replacement = mapping.get(f"{{{{{middle}}}}}")

            if replacement is None:
                # No es un placeholder conocido -> avanzamos un carácter para reintentar
                # (evita quedar pegados en el mismo '{{')
                # Partimos el run i en dos para “consumir” el '{{' hallado y seguir buscando.
                # Como python-docx no permite split directo, hacemos una sustitución temporal:
                pre = runs[i].text[:pos + 2]  # incluye '{{'
                post = runs[i].text[pos + 2:]
                runs[i].text = pre
                # Insertamos el post en el mismo run (concatenándolo) para no alterar estilos;
                # en la siguiente vuelta, no encontrará '{{' en este run y avanzará.
                runs[i].text = runs[i].text + post
                i += 1
                continue

            # Tenemos reemplazo. Reconstruimos: prefijo del primer run + replacement + sufijo global
            first_run_idx, from_pos, _ = consumed[0]
            # Prefijo: todo lo que hay antes de '{{' en el primer run
            prefix = runs[first_run_idx].text[:from_pos]

            # Sufijo: todo lo que viene después de '}}' a lo largo de los runs consumidos
            end_in_collected = close_idx + 2
            suffix = collected[end_in_collected:]  # puede estar vacío

            # Limpiamos todos los runs consumidos
            for run_idx, _, _ in consumed:
                runs[run_idx].text = ""

            # Escribimos en el primer run: prefix + replacement
            runs[first_run_idx].text = prefix + replacement

            # Si hay sufijo, lo colocamos en el último run consumido (mantiene el estilo de ese run)
            if suffix:
                last_run_idx = consumed[-1][0]
                runs[last_run_idx].text = suffix

            # No incrementamos demasiado 'i' para permitir detectar más placeholders seguidos
            i = max(i, first_run_idx)  # nos aseguramos de no retroceder
            i += 1

    def reemplazar_placeholder_word(self, archivo_docx: bytes, metadata: dict, formato: str) -> Union[str, bytes]:
        try:
            mapping = self._build_mapping(metadata)
            if not mapping:
                raise TributarioException("La metadata está vacía: no hay placeholders que reemplazar.")

            buffer_entrada = BytesIO(archivo_docx)
            doc = Document(buffer_entrada)

            # Párrafos
            for p in doc.paragraphs:
                self._replace_placeholders_in_runs(p.runs, mapping)

            # Tablas (ídem)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self._replace_placeholders_in_runs(p.runs, mapping)

            # Headers/Footers:
            for section in doc.sections:
                self._replace_placeholders_in_runs(section.header.paragraphs, mapping)
                self._replace_placeholders_in_runs(section.footer.paragraphs, mapping)
                for t in section.header.tables:
                    for r in t.rows:
                        for c in r.cells:
                            for p in c.paragraphs:
                                self._replace_placeholders_in_runs(p.runs, mapping)
                for t in section.footer.tables:
                    for r in t.rows:
                        for c in r.cells:
                            for p in c.paragraphs:
                                self._replace_placeholders_in_runs(p.runs, mapping)

            buffer_salida = BytesIO()
            doc.save(buffer_salida)
            buffer_salida.seek(0)
            contenido = buffer_salida.read()

            if formato == "base64":
                return base64.b64encode(contenido).decode("utf-8")
            elif formato == "file":
                return contenido
            else:
                raise TributarioException(
                    f"Se ingresó el formato '{formato}' que no es aceptado. Debe ser 'base64' o 'file'."
                )
        except TributarioException:
            raise
        except Exception as e:
            raise TributarioException(f"Error al procesar el Word: {str(e)}")
